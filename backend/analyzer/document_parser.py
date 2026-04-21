"""
Document Parser — Extracts text from PDF, DOCX, and TXT files.
Supports section-aware chunking for better semantic search.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional


def extract_text(file_path: Path) -> str:
    """Extract text from a file based on its extension."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    elif suffix == ".docx":
        return _extract_docx(file_path)
    elif suffix in (".txt", ".md"):
        return file_path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf(file_path: Path) -> str:
    """Extract text from PDF using PyPDF2 with pdfplumber fallback for tables."""
    try:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)

                # Extract tables
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            text_parts.append(
                                " | ".join(str(cell) for cell in row if cell)
                            )
        return "\n\n".join(text_parts)
    except ImportError:
        pass

    # Fallback to PyPDF2
    from PyPDF2 import PdfReader

    reader = PdfReader(str(file_path))
    text_parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text_parts.append(text)
    return "\n\n".join(text_parts)


def _extract_docx(file_path: Path) -> str:
    """Extract text from DOCX preserving paragraph structure."""
    from docx import Document

    doc = Document(str(file_path))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))

    return "\n\n".join(text_parts)


def chunk_text(
    text: str,
    max_chunk_size: int = 1000,
    overlap: int = 200,
) -> list[dict]:
    """
    Split text into overlapping chunks, trying to respect section boundaries.

    Returns list of dicts with 'text', 'index', and 'section' keys.
    """
    # Try to detect sections by headings
    sections = _split_by_sections(text)

    chunks = []
    chunk_index = 0

    for section_title, section_text in sections:
        # Split section into chunks if it's too long
        if len(section_text) <= max_chunk_size:
            chunks.append(
                {
                    "text": section_text,
                    "index": chunk_index,
                    "section": section_title,
                }
            )
            chunk_index += 1
        else:
            # Sliding window with overlap
            words = section_text.split()
            current_pos = 0
            while current_pos < len(words):
                end_pos = current_pos + max_chunk_size // 5  # ~5 chars per word
                chunk_words = words[current_pos:end_pos]
                chunk_text_piece = " ".join(chunk_words)

                chunks.append(
                    {
                        "text": chunk_text_piece,
                        "index": chunk_index,
                        "section": section_title,
                    }
                )
                chunk_index += 1

                # Move forward with overlap
                current_pos = max(current_pos + 1, end_pos - overlap // 5)

    return chunks


def _split_by_sections(text: str) -> list[tuple[str, str]]:
    """
    Split text into (section_title, section_content) tuples.
    Detects common heading patterns.
    """
    # Common heading patterns
    heading_patterns = [
        r"^#{1,4}\s+(.+)$",  # Markdown headings
        r"^(\d+\.[\d.]*\s+.+)$",  # Numbered sections (1. Introduction, 2.1 Data)
        r"^([A-Z][A-Z\s]{4,})$",  # ALL CAPS lines
    ]

    combined_pattern = "|".join(f"({p})" for p in heading_patterns)
    lines = text.split("\n")

    sections = []
    current_title = "Document Start"
    current_content: list[str] = []

    for line in lines:
        is_heading = False
        for pattern in heading_patterns:
            match = re.match(pattern, line.strip(), re.MULTILINE)
            if match:
                # Save previous section
                if current_content:
                    sections.append(
                        (current_title, "\n".join(current_content).strip())
                    )
                current_title = line.strip()
                current_content = []
                is_heading = True
                break

        if not is_heading:
            current_content.append(line)

    # Final section
    if current_content:
        sections.append((current_title, "\n".join(current_content).strip()))

    # If no sections detected, return entire text as one section
    if not sections:
        sections = [("Full Document", text)]

    return sections
